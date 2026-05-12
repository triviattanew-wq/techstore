from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_margins(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(3)
    s.right_margin = Cm(1.5)

def h(doc, text, size=14, bold=True, center=False):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(6)
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = 'Times New Roman'
    return par

def p(doc, text, size=12, bold=False, indent=True, center=False, italic=False):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.line_spacing = Pt(18)
    if indent:
        par.paragraph_format.first_line_indent = Cm(1.25)
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = 'Times New Roman'
    return par

def code(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(1)
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    return par

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, hd in enumerate(headers):
        hdr[i].text = hd
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    return table
